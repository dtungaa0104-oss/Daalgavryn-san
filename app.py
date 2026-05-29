import os, sqlite3, re
from datetime import datetime
from functools import wraps
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
try:
    import anthropic
except ImportError:
    anthropic = None
try:
    from blueprint_data import get_blueprint_route, save_blueprint_route, init_blueprint_table
    _HAS_BLUEPRINT = True
except ImportError:
    _HAS_BLUEPRINT = False
# Curriculum data
import json as _json
_curr_path = os.path.join(os.path.dirname(__file__), 'templates', 'curriculum.json')
try:
    with open(_curr_path, 'r', encoding='utf-8') as _f:
        CURRICULUM = _json.load(_f)
except:
    CURRICULUM = {"grades": {}, "periods": ["40","80","90"]}

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "orkhontul-ebs-2025")
app.config.update(
    SESSION_COOKIE_SECURE   = False,
    SESSION_COOKIE_HTTPONLY = True,
    SESSION_COOKIE_SAMESITE = "Lax",
    SESSION_COOKIE_NAME     = "daalgavar_session",
    PERMANENT_SESSION_LIFETIME = 86400 * 365,
    MAX_CONTENT_LENGTH      = 50 * 1024 * 1024  # 50MB
)
# Render HTTPS proxy дамжуулах
try:
    from werkzeug.middleware.proxy_fix import ProxyFix
    app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)
except Exception:
    pass
# ── DATABASE: PostgreSQL (Supabase) эсвэл SQLite fallback ──────────────
DATABASE_URL     = os.environ.get("DATABASE_URL", "postgresql://postgres.vnxqgqthvqhyziwyyvsm:Orkhontuul%402020@aws-0-ap-southeast-1.pooler.supabase.com:6543/postgres")
SUPABASE_URL     = os.environ.get("SUPABASE_URL", "https://vnxqgqthvqhyziwyyvsm.supabase.co")
SUPABASE_KEY     = os.environ.get("SUPABASE_KEY", "")  # service_role key
DB_PATH      = os.environ.get("DB_PATH", "questions.db")
USE_PG       = bool(DATABASE_URL)

try:
    import psycopg2
    import psycopg2.extras
    _HAS_PG = True
    print("✅ psycopg2 байна — PostgreSQL ашиглана")
except ImportError:
    _HAS_PG = False
    psycopg2 = None

# DB сонгох — SUPABASE_KEY байвал REST API, эсвэл SQLite
if SUPABASE_KEY:
    USE_PG = True
    print("✅ Supabase REST API ашиглана")
elif _HAS_PG and DATABASE_URL:
    USE_PG = True
    print("✅ PostgreSQL (psycopg2) ашиглана")
else:
    USE_PG = False
    print("⚠️ SQLite ашиглана (data устна!) — SUPABASE_KEY тавина уу")

EXAM_TYPES = {
    "ulsiin": {
        "id":"ulsiin","name":"Улсын шалгалт","icon":"🏆","color":"#f0a500",
        "grades":[5,9,12],
        "blueprint":{
            "total":20,"score":35,"duration":"80 минут",
            "Мэдлэг ойлголт":6,"Чадвар":8,"Хэрэглээ":6,
            "zadgai_total":3,"zadgai_score":9,
            "zadgai": [
                {"num":"2.1","score":3,"desc":"Харьцаа, пропорц"},
                {"num":"2.2","score":3,"desc":"Геометр, биет"},
                {"num":"2.3","score":3,"desc":"Статистик, практик"},
            ],
            "note":"2025: 17 сонгох (1-2 оноо) + 9 задгай = 35 оноо · 80 минут"
        }
    },
    "devshih": {
        "id":"devshih","name":"Анги дэвших шалгалт","icon":"📋","color":"#3b82f6",
        "grades":list(range(3,12)),
        "blueprint":{
            "total":25,"score":49,"duration":"60 минут",
            "Мэдлэг ойлголт":8,"Чадвар":10,"Хэрэглээ":7,
            "zadgai_total":0,"zadgai_score":0,"zadgai":[],
            "note":"25 сонгох даалгавар · 60 минут"
        }
    },
    "guitsegdel": {
        "id":"guitsegdel","name":"Гүйцэтгэлийн үнэлгээ","icon":"📝","color":"#22c55e",
        "grades":list(range(3,13)),
        "blueprint":{
            "total":10,"score":0,"duration":"30 минут",
            "Мэдлэг ойлголт":4,"Чадвар":4,"Хэрэглээ":2,
            "zadgai_total":0,"zadgai_score":0,"zadgai":[],
            "note":"Оноогоор дүгнэхгүй — гүйцэтгэлийн үнэлгээ"
        }
    },
    "elselt": {
        "id":"elselt","name":"Элсэлтийн шалгалт","icon":"🎓","color":"#a855f7",
        "grades":list(range(11,13)),
        "blueprint":{
            "total":40,"score":100,"duration":"90 минут",
            "Мэдлэг ойлголт":10,"Чадвар":18,"Хэрэглээ":12,
            "zadgai_total":0,"zadgai_score":0,"zadgai":[],
            "note":"1 буруу = −0.2 оноо · Нийт 100 оноо"
        }
    },
}

GRADE_EXAM_MAP = {}
for eid, et in EXAM_TYPES.items():
    for g in et["grades"]:
        if g not in GRADE_EXAM_MAP:
            GRADE_EXAM_MAP[g] = eid
GRADE_EXAM_MAP[1] = "guitsegdel"

SUBJECTS = {
    "1-5":   ["Монгол хэл, уран зохиол","Математик","Байгалийн ухаан","Нийгэм судлал"],
    "6-9":   ["Монгол хэл, уран зохиол","Физик","Хими","Биологи","Газарзүй","Түүх","Англи хэл"],
    "10-12": ["Монгол хэл, уран зохиол","Математик","Физик","Хими","Биологи","Газарзүй","Түүх","Англи хэл","Нийгмийн ухаан"]
}

ALL_SUBJECTS = [
    "Монгол хэл, уран зохиол","Монгол хэл","Математик","Байгалийн ухаан","Нийгэм судлал",
    "Физик","Хими","Биологи","Газарзүй","Түүх","Англи хэл","Орос хэл","Нийгмийн ухаан",
    "Мэдээллийн технологи","Хөгжим","Технологи","Дүрслэх урлаг",
    "Биеийн тамир","Иргэний ёс зүйн боловсрол","Эрүүл мэнд"
]

LEVELS      = ["Мэдлэг ойлголт","Чадвар","Хэрэглээ"]
BLOOM       = ["Мэдлэг","Ойлголт","Хэрэглээ","Шинжилгээ","Үнэлгээ","Бүтээл"]
Q_TYPES     = ["Нэг сонголт","Олон сонголт","Нээлттэй","Гүйцэтгэлийн","Олимпиад"]
ADMIN_PW    = os.environ.get("ADMIN_PASSWORD","orkhontul2025")
LEVEL_SCORE = {"Мэдлэг ойлголт":1,"Чадвар":2,"Хэрэглээ":3}
TEACHER_PASSWORD = os.environ.get("TEACHER_PASSWORD", "orkhontul-bagsh")

class SupabaseRestConn:
    """Supabase REST API (PostgREST) ашиглан CRUD хийх — psycopg2 шаардахгүй"""
    def __init__(self):
        self._rows      = []
        self._lastrowid = None
    def _hdr(self):
        return {
            "apikey":        SUPABASE_KEY,
            "Authorization": "Bearer " + SUPABASE_KEY,
            "Content-Type":  "application/json",
            "Prefer":        "return=representation"
        }
    def _url(self, table):
        return SUPABASE_URL + "/rest/v1/" + table
    def execute(self, sql, params=()):
        import re as _re, requests as _rq, json as _j
        sql_s = sql.strip()
        sql_u = sql_s.upper().replace("\n", " ").replace("  ", " ")
        try:
            if "SELECT COUNT(*)" in sql_u:
                m = _re.search(r"FROM\s+(\w+)(.*)", sql_s, _re.IGNORECASE | _re.DOTALL)
                if not m:
                    self._rows = [{"count": 0}]; return self
                table = m.group(1).strip()
                rest  = m.group(2).strip() if m.group(2) else ""
                hdrs  = dict(self._hdr())
                hdrs["Prefer"] = "count=exact"
                params_url = "?select=id"
                if "WHERE" in rest.upper():
                    params_url += self._where_params(rest, params)
                r = _rq.head(self._url(table) + params_url, headers=hdrs, timeout=10)
                cr = r.headers.get("content-range", "0/0")
                cnt = int(cr.split("/")[-1]) if cr and "/" in cr else 0
                self._rows = [{"count": cnt}]
                return self
            elif sql_u.startswith("SELECT"):
                m = _re.search(r"FROM\s+(\w+)(.*)", sql_s, _re.IGNORECASE | _re.DOTALL)
                if not m:
                    self._rows = []; return self
                table = m.group(1).strip()
                rest  = m.group(2).strip() if m.group(2) else ""
                url   = self._url(table) + "?select=*"
                if "WHERE" in rest.upper():
                    url += self._where_params(rest, params)
                m2 = _re.search(r"LIMIT\s+(\d+)", sql_u)
                if m2: url += "&limit=" + m2.group(1)
                url += "&limit=1000"
                r = _rq.get(url, headers=self._hdr(), timeout=15)
                self._rows = r.json() if r.status_code == 200 else []
                return self
            elif sql_u.startswith("INSERT"):
                m = _re.search(r"INSERT\s+(?:OR\s+IGNORE\s+|OR\s+REPLACE\s+)?INTO\s+(\w+)\s*\(([^)]+)\)", sql_s, _re.IGNORECASE | _re.DOTALL)
                if not m:
                    print(f"SupabaseREST INSERT parse fail: {sql_s[:80]}")
                    return self
                table = m.group(1).strip()
                cols  = [col.strip() for col in m.group(2).split(",")]
                row   = {}
                for i, col in enumerate(cols):
                    if i < len(params) and params[i] is not None:
                        val = params[i]
                        # nj багана JSONB — JSON string-г parse хийж object болгох
                        if col == "nj" and isinstance(val, str):
                            try:
                                val = _j.loads(val)
                            except Exception:
                                pass
                        row[col] = val
                hdrs = dict(self._hdr())
                hdrs["Prefer"] = "return=representation,resolution=ignore-duplicates"
                r = _rq.post(self._url(table), headers=hdrs, json=row, timeout=15)
                if r.status_code in (200, 201):
                    data = r.json()
                    if data and isinstance(data, list) and data[0].get("id"):
                        self._lastrowid = data[0]["id"]
                elif r.status_code == 409:
                    pass
                else:
                    print(f"SupabaseREST INSERT {table} → {r.status_code}: {r.text[:200]}")
                return self
            elif sql_u.startswith("DELETE"):
                m = _re.search(r"FROM\s+(\w+)\s+WHERE\s+id\s*=", sql_s, _re.IGNORECASE)
                if m and params:
                    table = m.group(1).strip()
                    _rq.delete(self._url(table) + "?id=eq." + str(params[0]), headers=self._hdr(), timeout=10)
                    return self
                # WHERE subject=? AND grade=? (blueprints)
                m2 = _re.search(r"FROM\s+(\w+)\s+WHERE\s+(.+)", sql_s, _re.IGNORECASE | _re.DOTALL)
                if m2 and params:
                    table = m2.group(1).strip()
                    qp = self._where_params("WHERE " + m2.group(2), params)
                    if qp:
                        _rq.delete(self._url(table) + "?" + qp.lstrip("&"), headers=self._hdr(), timeout=10)
                return self
            elif sql_u.startswith("UPDATE"):
                m = _re.search(r"UPDATE\s+(\w+)\s+SET\s+(.+?)\s+WHERE\s+id\s*=", sql_s, _re.IGNORECASE | _re.DOTALL)
                if m and params:
                    table = m.group(1).strip()
                    set_part = m.group(2)
                    cols = [c2.split("=")[0].strip() for c2 in set_part.split(",")]
                    update_data = {col: params[i] for i, col in enumerate(cols) if i < len(params)-1}
                    _rq.patch(self._url(table) + "?id=eq." + str(params[-1]), headers=self._hdr(), json=update_data, timeout=10)
                return self
            elif sql_u.startswith(("CREATE", "ALTER", "DROP")):
                return self
        except Exception as e:
            print(f"SupabaseREST error: {e}")
        return self
    def _where_params(self, rest, params):
        import re as _re
        result = ""
        wi = rest.upper().find("WHERE")
        if wi < 0: return result
        where = rest[wi+5:].strip()
        parts = _re.split(r"\s+AND\s+", where, flags=_re.IGNORECASE)
        pi = 0
        for part in parts:
            part = part.strip()
            m = _re.match(r"(\w+)\s*=\s*(?:[?]|%s)", part, _re.IGNORECASE)
            if m and pi < len(params):
                col = m.group(1)
                val = params[pi]
                result += f"&{col}=eq.{val}"
                pi += 1
        return result
    def fetchall(self):
        return self._rows or []
    def fetchone(self):
        rows = self._rows or []
        return rows[0] if rows else None
    @property
    def lastrowid(self):
        return self._lastrowid
    def commit(self): pass
    def close(self):
        try:
            self._conn.close()
        except Exception:
            pass


class PGConn:
    """psycopg2 connection-г sqlite3-тай адил interface болгох wrapper"""
    def __init__(self, conn):
        self._conn = conn
    def execute(self, sql, params=()):
        sql_pg = sql.replace("?", "%s")
        sql_pg = sql_pg.replace("INSERT OR IGNORE INTO", "INSERT INTO")
        sql_pg = sql_pg.replace("INSERT OR REPLACE INTO", "INSERT INTO")
        is_insert = sql_pg.strip().upper().startswith("INSERT")
        if is_insert and "RETURNING" not in sql_pg.upper() and "ON CONFLICT" not in sql_pg.upper():
            last = sql_pg.rfind(")")
            if last >= 0:
                sql_pg = sql_pg[:last+1] + " ON CONFLICT (q_code) DO NOTHING RETURNING id"
        cur = self._conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql_pg, params)
        if is_insert:
            try:
                row = cur.fetchone()
                if row and 'id' in row:
                    cur.lastrowid = row['id']
                else:
                    cur.lastrowid = None
            except:
                cur.lastrowid = None
        else:
            cur.lastrowid = None
        return cur
    def executemany(self, sql, params_list):
        sql_pg = sql.replace("?", "%s")
        cur = self._conn.cursor()
        cur.executemany(sql_pg, params_list)
        return cur
    def commit(self):
        self._conn.commit()
    def close(self):
        self._conn.close()
    def __enter__(self):
        return self
    def __exit__(self, *a):
        self.close()


def get_db():
    """PostgreSQL (psycopg2) / Supabase REST / SQLite буцаана"""
    if USE_PG and _HAS_PG:
        try:
            conn = psycopg2.connect(DATABASE_URL, sslmode='require', connect_timeout=10)
            return PGConn(conn)
        except Exception as e:
            print(f"PG холболт алдаа: {e}")
            # psycopg2 амжилтгүй — SUPABASE_KEY байвал REST руу унана
            if SUPABASE_KEY:
                print(" → Supabase REST API ашиглана")
                return SupabaseRestConn()
            print(" → SQLite ашиглана")
    elif USE_PG and SUPABASE_KEY and not _HAS_PG:
        return SupabaseRestConn()
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def _fetch_scalar(conn, sql, params=()):
    try:
        row = conn.execute(sql, params).fetchone()
        if row is None: return 0
        if isinstance(row, dict): return list(row.values())[0]
        try: return row[0]
        except: return list(dict(row).values())[0]
    except Exception:
        return 0


def init_db():
    import sqlite3 as _sq
    conn = get_db()
    is_pg = USE_PG and _HAS_PG and not isinstance(conn, _sq.Connection)
    if is_pg:
        sql_q = """CREATE TABLE IF NOT EXISTS questions (
            id SERIAL PRIMARY KEY, q_code TEXT UNIQUE,
            grade INTEGER NOT NULL, subject TEXT NOT NULL,
            level TEXT NOT NULL DEFAULT 'Мэдлэг ойлголт',
            bloom TEXT NOT NULL DEFAULT 'Мэдлэг',
            q_type TEXT NOT NULL DEFAULT 'Нэг сонголт',
            question TEXT NOT NULL,
            option_a TEXT, option_b TEXT, option_c TEXT, option_d TEXT,
            answer TEXT, score INTEGER DEFAULT 1,
            topic TEXT DEFAULT '', hint TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""
        sql_t = """CREATE TABLE IF NOT EXISTS question_topics (
            id SERIAL PRIMARY KEY, grade INTEGER,
            subject TEXT, topic TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""
    else:
        sql_q = """CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT, q_code TEXT UNIQUE,
            grade INTEGER NOT NULL, subject TEXT NOT NULL,
            level TEXT NOT NULL DEFAULT 'Мэдлэг ойлголт',
            bloom TEXT NOT NULL DEFAULT 'Мэдлэг',
            q_type TEXT NOT NULL DEFAULT 'Нэг сонголт',
            question TEXT NOT NULL,
            option_a TEXT, option_b TEXT, option_c TEXT, option_d TEXT,
            answer TEXT, score INTEGER DEFAULT 1,
            topic TEXT DEFAULT '', hint TEXT DEFAULT '',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP)"""
        sql_t = """CREATE TABLE IF NOT EXISTS question_topics (
            id INTEGER PRIMARY KEY AUTOINCREMENT, grade INTEGER,
            subject TEXT, topic TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP)"""
    conn.execute(sql_q)
    conn.execute(sql_t)
    conn.commit()
    conn.close()


def row_to_dict(row):
    d = dict(row)
    d["options"] = (
        [f"А. {d['option_a']}",f"Б. {d['option_b']}",
         f"В. {d['option_c']}",f"Г. {d['option_d']}"]
        if d.get("option_a") else None)
    return d

def login_required(f):
    @wraps(f)
    def dec(*a,**kw):
        if not session.get("admin"):
            return redirect(url_for("admin_login"))
        return f(*a,**kw)
    return dec

def teacher_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get("teacher_logged_in") or session.get("admin"):
            return f(*args, **kwargs)
        return redirect(url_for("teacher_login_page"))
    return decorated


def _get_blueprint_list():
    try:
        conn = get_db()
        rows = conn.execute(
            "SELECT subject, grade, nj FROM blueprints ORDER BY subject, grade"
        ).fetchall()
        conn.close()
        import json as _j
        result = []
        for r in rows:
            nj = r['nj'] if isinstance(r['nj'], list) else _j.loads(r.get('nj') or '[]')
            result.append({"subject": r['subject'], "grade": r['grade'], "nj_count": len(nj)})
        return result
    except Exception:
        return []

# ══ PUBLIC ════════════════════════════════════════════════
@app.route("/")
def index():
    return render_template("index.html", subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS, exam_types=EXAM_TYPES, grade_exam_map=GRADE_EXAM_MAP)

@app.route("/questions")
def questions_page():
    return render_template("questions.html", subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS, levels=LEVELS, blooms=BLOOM,
        exam_types=EXAM_TYPES, grade_exam_map=GRADE_EXAM_MAP)

@app.route("/blueprint")
def blueprint_page():
    return render_template("blueprint.html", subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS, exam_types=EXAM_TYPES, levels=LEVELS,
        blooms=BLOOM, grade_exam_map=GRADE_EXAM_MAP)

@app.route("/olympiad")
def olympiad_page():
    return render_template("olympiad.html", subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS, levels=LEVELS, blooms=BLOOM, grade_exam_map=GRADE_EXAM_MAP)

@app.route("/interactive")
def interactive_page():
    return render_template("interactive.html", subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS, levels=LEVELS, blooms=BLOOM, grade_exam_map=GRADE_EXAM_MAP)

# ══ БЛЮПРИНТ API ══════════════════════════════════════════
@app.route("/api/get-blueprint", methods=["GET"])
def api_get_blueprint():
    if not _HAS_BLUEPRINT:
        return jsonify({"blueprint": [], "warning": "blueprint_data.py олдсонгүй"}), 200
    return get_blueprint_route()

@app.route("/api/save-blueprint", methods=["POST"])
def api_save_blueprint():
    if not _HAS_BLUEPRINT:
        return jsonify({"error": "blueprint_data.py олдсонгүй"}), 500
    return save_blueprint_route()

# ══ БОСГО ШАЛГАЛТ ═════════════════════════════════════════
@app.route("/mongol-bichig")
def mongol_bichig_page():
    conn = get_db()
    years = conn.execute(
        "SELECT DISTINCT topic FROM questions WHERE subject='Монгол хэл бичгийн босго шалгалт' ORDER BY topic DESC"
    ).fetchall()
    recent = conn.execute(
        "SELECT * FROM questions WHERE subject='Монгол хэл бичгийн босго шалгалт' ORDER BY id DESC LIMIT 20"
    ).fetchall()
    total = _fetch_scalar(conn,
        "SELECT COUNT(*) FROM questions WHERE subject='Монгол хэл бичгийн босго шалгалт'")
    conn.close()
    return render_template("mongol_bichig.html",
        years=[r['topic'] for r in years], recent=recent, total=total,
        levels=LEVELS, blooms=BLOOM)

@app.route("/api/update-topic")
def update_topic():
    subject = request.args.get("subject","")
    topic   = request.args.get("topic","")
    if not subject or not topic:
        return jsonify({"ok": False})
    conn = get_db()
    conn.execute("UPDATE questions SET topic=? WHERE subject=? AND (topic IS NULL OR topic='')", (topic, subject))
    conn.commit(); conn.close()
    return jsonify({"ok": True})

# ══ API ═══════════════════════════════════════════════════
@app.route("/api/questions")
def api_questions():
    grade   = request.args.get("grade","")
    subject = request.args.get("subject","")
    level   = request.args.get("level","all")
    bloom   = request.args.get("bloom","all")
    q_type  = request.args.get("q_type","all")
    count   = int(request.args.get("count",20))
    conn=get_db(); sql,params="SELECT * FROM questions WHERE 1=1",[]
    if grade:         sql+=" AND grade=?";   params.append(int(grade))
    if subject:       sql+=" AND subject=?"; params.append(subject)
    if level!="all":  sql+=" AND level=?";   params.append(level)
    if bloom!="all":  sql+=" AND bloom=?";   params.append(bloom)
    if q_type!="all": sql+=" AND q_type=?";  params.append(q_type)
    sql+=" ORDER BY RANDOM() LIMIT ?"; params.append(count)
    rows=conn.execute(sql,params).fetchall(); conn.close()
    return jsonify({"questions":[row_to_dict(r) for r in rows],"total":len(rows)})


@app.route("/api/export-docx", methods=["POST"])
def export_docx():
    import io as _io
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    data     = request.json
    title    = data.get("title", "Шалгалт")
    grade    = data.get("grade", "")
    subject  = data.get("subject", "")
    duration = data.get("duration", "")
    questions= data.get("questions", [])
    show_ans = data.get("show_answers", False)
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2)
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.runs[0]; run.font.size = Pt(14); run.font.bold = True
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{grade}-р анги  ·  {subject}  ·  {duration}").font.size = Pt(11)
    doc.add_paragraph("─" * 60)
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        num_run = p.add_run(f"{i}. "); num_run.bold = True; num_run.font.size = Pt(11)
        q_run = p.add_run(q.get("question", "")); q_run.font.size = Pt(11)
        score_run = p.add_run(f"  /{q.get('score',1)} оноо/"); score_run.font.size = Pt(9)
        score_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.paragraph_format.space_before = Pt(6)
        opts = q.get("options", [])
        if opts:
            labels = ["А", "Б", "В", "Г"]
            opt_p = doc.add_paragraph()
            opt_p.paragraph_format.left_indent = Cm(1)
            opt_p.paragraph_format.space_before = Pt(2)
            for j, opt in enumerate(opts):
                label = labels[j] if j < len(labels) else str(j+1)
                is_ans = show_ans and label == q.get("answer","")
                run = opt_p.add_run(f"{label}. {opt}     "); run.font.size = Pt(10)
                if is_ans:
                    run.bold = True; run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    if show_ans:
        doc.add_page_break()
        doc.add_heading("Зөв хариулт", level=1)
        ans_p = doc.add_paragraph()
        for i, q in enumerate(questions, 1):
            ans_p.add_run(f"{i}. {q.get('answer','')}   ")
    buf = _io.BytesIO(); doc.save(buf); buf.seek(0)
    from flask import Response
    safe_name = title.replace(" ", "_").replace("/","")[:50]
    return Response(buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_name}.docx"})

@app.route("/api/generate-exam",methods=["POST"])
def generate_exam():
    data=request.json; grade=data.get("grade",9); subject=data.get("subject","Математик")
    exam_id=data.get("exam_id","devshih"); blueprint=data.get("blueprint",{})
    use_ai=data.get("use_ai",True)
    et=EXAM_TYPES.get(exam_id,EXAM_TYPES["devshih"]); bp=et["blueprint"]
    conn=get_db(); selected=[]; ai_generated=0
    for lvl in LEVELS:
        cnt=int(blueprint.get(lvl,bp.get(lvl,0)))
        if cnt<=0: continue
        rows=conn.execute(
            "SELECT * FROM questions WHERE grade=? AND subject=? AND level=? ORDER BY RANDOM() LIMIT ?",
            (grade,subject,lvl,cnt)).fetchall()
        db_qs=[row_to_dict(r) for r in rows]
        selected.extend(db_qs)
        need=cnt-len(db_qs)
        if need>0 and use_ai:
            api_key=os.environ.get("ANTHROPIC_API_KEY","")
            if api_key:
                try:
                    import json as _json
                    score=LEVEL_SCORE.get(lvl,1)
                    lvl_guide = {'Мэдлэг ойлголт':'Тодорхойлолт таних, томьёо мэдэх, тооцоолол','Чадвар':'Бодлого шийдэх, тайлбарлах, хэрэглэх','Хэрэглээ':'Амьдралын бодлого, дүгнэх, загвар байгуулах'}
                    is_math = subject in ['Математик','Алгебр','Геометр','Тригонометр']
                    math_hint = (
                        'Томьёог ЗААВАЛ LaTeX форматаар бич ($...$). '
                        'Жишээ: зэрэг=$a^2+b^2=c^2$, язгуур=$\\sqrt{x+1}$, '
                        'бутархай=$\\frac{a}{b}$, тригонометр=$\\sin^2x+\\cos^2x=1$. '
                        'Геометрийн бодлогод: гурвалжин/тойрог/куб/цилиндр гэсэн үг заавал оруул.'
                        if is_math else '')
                    prompt = (
                        f'Монгол ЕБС {grade}-р анги {subject} хичээл. '
                        f'Блюпринт: {lvl} ({lvl_guide.get(lvl,"")}).{math_hint}\n'
                        f'{need} даалгавар үүсгэ. Монгол хэлээр. '
                        'Зөвхөн JSON array буцаа, тайлбаргүй:\n'
                        '[{"question":"...","option_a":"...","option_b":"...","option_c":"...","option_d":"...","answer":"А"}]'
                    )
                    client=anthropic.Anthropic(api_key=api_key)
                    msg=client.messages.create(model="claude-sonnet-4-5",max_tokens=3000,
                       messages=[{"role":"user","content":prompt}])
                    raw=msg.content[0].text.strip()
                    raw=re.sub(r'^```json\s*','',raw); raw=re.sub(r'^```\s*','',raw); raw=re.sub(r'\s*```$','',raw).strip()
                    ai_qs=_json.loads(raw)
                    if not isinstance(ai_qs,list): ai_qs=[ai_qs]
                    for idx,q in enumerate(ai_qs[:need]):
                       q_code=f"Q{grade}-{subject[:2]}-AI-{datetime.now().strftime('%f')}-{idx}"
                       try:
                           conn.execute("""INSERT INTO questions(q_code,grade,subject,level,bloom,q_type,
                               question,option_a,option_b,option_c,option_d,answer,score,topic)
                               VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                               (q_code,grade,subject,lvl,'Мэдлэг','Нэг сонголт',q.get('question',''),
                                q.get('option_a'),q.get('option_b'),q.get('option_c'),q.get('option_d'),
                                q.get('answer','А'),score,''))
                           conn.commit()
                           saved=conn.execute("SELECT * FROM questions WHERE q_code=?",(q_code,)).fetchone()
                           if saved:
                               selected.append(row_to_dict(saved))
                               ai_generated+=1
                       except: pass
                except Exception as ai_err:
                    print(f"AI generate error: {ai_err}")
    conn.close()
    if not selected:
        api_key=os.environ.get("ANTHROPIC_API_KEY","")
        if not api_key:
            return jsonify({"error":f"{grade}-р ангийн '{subject}' даалгавар байхгүй. ANTHROPIC_API_KEY тохируулна уу."})
        return jsonify({"error":f"{grade}-р ангийн '{subject}' даалгавар үүсгэж чадсангүй."})
    msg_parts=[]
    db_count=len(selected)-ai_generated
    if db_count>0: msg_parts.append(f"DB: {db_count}")
    if ai_generated>0: msg_parts.append(f"AI: {ai_generated}")
    return jsonify({"exam_id":exam_id,"exam_type":et["name"],"exam_icon":et["icon"],
        "title":f"{grade}-р ангийн {subject} — {et['name']}","grade":grade,"subject":subject,
        "total_questions":len(selected),"total_score":sum(q["score"] for q in selected),
        "duration":bp["duration"],"note":bp.get("note",""),
        "ai_generated":ai_generated,
        "source_note":" · ".join(msg_parts) if msg_parts else "",
        "blueprint_info":{
            "zadgai_total": bp.get("zadgai_total",0),
            "zadgai_score": bp.get("zadgai_score",0),
            "zadgai": bp.get("zadgai",[]),
        },
        "questions":selected})


@app.route("/api/db-init")
def db_init():
    try:
        init_db()
        conn = get_db()
        cnt = _fetch_scalar(conn, "SELECT COUNT(*) FROM questions") or 0
        conn.close()
        return jsonify({"status": "ok",
            "db": "rest_api" if (USE_PG and not _HAS_PG and SUPABASE_KEY) else
                  "postgresql" if (USE_PG and _HAS_PG) else "sqlite",
            "questions": cnt})
    except Exception as e:
        import traceback
        return jsonify({"status": "error", "error": str(e), "trace": traceback.format_exc()[:500]}), 500

@app.route("/api/version")
def version():
    return jsonify({"version": "2025-v6",
        "db": "postgresql" if (USE_PG and _HAS_PG) else "sqlite",
        "psycopg2": _HAS_PG, "use_pg": USE_PG})

@app.route("/api/stats")
def stats():
    conn=get_db()
    total=_fetch_scalar(conn, "SELECT COUNT(*) FROM questions")
    grades=_fetch_scalar(conn, "SELECT COUNT(DISTINCT grade) FROM questions")
    subjects=_fetch_scalar(conn, "SELECT COUNT(DISTINCT subject) FROM questions")
    conn.close()
    return jsonify({"total_questions":total,"grades":grades or 12,
                    "subjects":subjects or 10,"blueprints":len(EXAM_TYPES)})

# ══ ADMIN ═════════════════════════════════════════════════
@app.route("/admin/login",methods=["GET","POST"])
def admin_login():
    error=None
    if request.method=="POST":
        if request.form.get("password")==ADMIN_PW:
            session.permanent = True
            session["admin"] = True
            return redirect(url_for("admin_dashboard"))
        error="Нууц үг буруу байна!"
    return render_template("admin_login.html",error=error)

@app.route("/admin/logout")
def admin_logout():
    session.pop("admin",None); return redirect(url_for("admin_login"))

@app.route("/admin")
@login_required
def admin_dashboard():
    conn=get_db()
    total=_fetch_scalar(conn, "SELECT COUNT(*) FROM questions")
    by_level=conn.execute("SELECT level,COUNT(*) cnt FROM questions GROUP BY level").fetchall()
    by_sub=conn.execute("SELECT subject,COUNT(*) cnt FROM questions GROUP BY subject ORDER BY cnt DESC").fetchall()
    recent=conn.execute("SELECT * FROM questions ORDER BY id DESC LIMIT 8").fetchall()
    conn.close()
    return render_template("admin_dashboard.html",total=total,by_level=by_level,
        by_sub=by_sub,recent=recent,subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS,levels=LEVELS,blooms=BLOOM,q_types=Q_TYPES)

@app.route("/admin/add",methods=["GET","POST"])
@login_required
def admin_add():
    if request.method=="POST":
        f=request.form; lvl=f["level"]; score=LEVEL_SCORE.get(lvl,1)
        q_code=f"Q{f['grade']}-{f['subject'][:3]}-{datetime.now().strftime('%f')}"
        conn=get_db()
        try:
            conn.execute("""INSERT INTO questions(q_code,grade,subject,level,bloom,q_type,question,
                option_a,option_b,option_c,option_d,answer,score,topic)VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (q_code,int(f["grade"]),f["subject"],lvl,f["bloom"],f["q_type"],f["question"],
                 f.get("option_a"),f.get("option_b"),f.get("option_c"),f.get("option_d"),
                 f.get("answer"),score,f.get("topic","")))
            conn.commit()
        except Exception as _e:
            print(f"admin_add INSERT error: {_e}")
            import traceback; traceback.print_exc()
        conn.close()
        return redirect(url_for("admin_list"))
    return render_template("admin_add.html",q=None,
        subjects=SUBJECTS, all_subjects=ALL_SUBJECTS,levels=LEVELS,blooms=BLOOM,q_types=Q_TYPES)

@app.route("/admin/edit/<int:qid>",methods=["GET","POST"])
@login_required
def admin_edit(qid):
    conn=get_db()
    if request.method=="POST":
        f=request.form; lvl=f["level"]; score=LEVEL_SCORE.get(lvl,1)
        conn.execute("""UPDATE questions SET grade=?,subject=?,level=?,bloom=?,q_type=?,
            question=?,option_a=?,option_b=?,option_c=?,option_d=?,answer=?,score=?,topic=? WHERE id=?""",
            (int(f["grade"]),f["subject"],lvl,f["bloom"],f["q_type"],f["question"],
             f.get("option_a"),f.get("option_b"),f.get("option_c"),f.get("option_d"),
             f.get("answer"),score,f.get("topic",""),qid))
        conn.commit(); conn.close(); return redirect(url_for("admin_list"))
    q=conn.execute("SELECT * FROM questions WHERE id=?",(qid,)).fetchone(); conn.close()
    return render_template("admin_add.html",q=q,
        subjects=SUBJECTS, all_subjects=ALL_SUBJECTS,levels=LEVELS,blooms=BLOOM,q_types=Q_TYPES)

@app.route("/admin/list")
@login_required
def admin_list():
    grade=request.args.get("grade",""); subject=request.args.get("subject","")
    lvl=request.args.get("level","")
    conn=get_db(); sql,params="SELECT * FROM questions WHERE 1=1",[]
    if grade:   sql+=" AND grade=?";   params.append(int(grade))
    if subject: sql+=" AND subject=?"; params.append(subject)
    if lvl:     sql+=" AND level=?";   params.append(lvl)
    sql+=" ORDER BY id DESC"
    rows=conn.execute(sql,params).fetchall(); conn.close()
    return render_template("admin_list.html",questions=rows,subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS,levels=LEVELS,
        sel_grade=grade,sel_subject=subject,sel_level=lvl)

@app.route("/admin/delete/<int:qid>",methods=["POST"])
@login_required
def admin_delete(qid):
    conn=get_db(); conn.execute("DELETE FROM questions WHERE id=?",(qid,))
    conn.commit(); conn.close(); return redirect(url_for("admin_list"))

@app.route("/admin/import",methods=["GET","POST"])
@login_required
def admin_import():
    if request.method=="POST":
        from file_importer import extract_from_pdf,extract_from_docx,parse_raw_text
        f=request.files.get("file"); grade=int(request.form.get("grade",9))
        subject=request.form.get("subject","Математик")
        lvl=request.form.get("default_level","auto")
        if not f or f.filename=="":
            return jsonify({"error":"Файл сонгоогүй байна"}),400
        file_bytes=f.read(); fname=f.filename.lower()
        try:
            if fname.endswith(".pdf"):
                import pdfplumber, io as _io
                with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                    pages = pdf.pages[:50]
                    raw = "\n".join(p.extract_text() or "" for p in pages)
            elif fname.endswith((".docx",".doc")): raw=extract_from_docx(file_bytes)
            elif fname.endswith(".txt"):            raw=file_bytes.decode("utf-8",errors="ignore")
            elif fname.endswith((".jpg",".jpeg",".png")):
                import base64, anthropic
                b64 = base64.b64encode(file_bytes).decode()
                mt = "image/jpeg" if fname.endswith((".jpg",".jpeg")) else "image/png"
                api_key2=os.environ.get("ANTHROPIC_API_KEY","")
                if not api_key2:
                    return jsonify({"error":"Зураг задлахад ANTHROPIC_API_KEY шаардлагатай"}),400
                cl2=anthropic.Anthropic(api_key=api_key2)
                msg2=cl2.messages.create(model="claude-sonnet-4-5",max_tokens=2000,
                    messages=[{"role":"user","content":[
                       {"type":"image","source":{"type":"base64","media_type":mt,"data":b64}},
                       {"type":"text","text":"Зурган дээрх даалгаваруудыг дугаарлан задла."}
                    ]}])
                raw=msg2.content[0].text
            else: return jsonify({"error":"PDF, Word, JPG эсвэл TXT файл оруулна уу"}),400
            questions=parse_raw_text(raw,grade,subject,lvl)
            if not questions:
                return jsonify({"error":"Даалгавар илрүүлж чадсангүй."}),400
            init_db()
            conn=get_db(); saved=skipped=0
            for q in questions:
                try:
                    conn.execute("""INSERT OR IGNORE INTO questions(q_code,grade,subject,level,bloom,q_type,
                       question,option_a,option_b,option_c,option_d,answer,score,topic)
                       VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                       (q['q_code'],q['grade'],q['subject'],q.get('level',lvl),q['bloom'],q['q_type'],
                        q['question'],q['option_a'],q['option_b'],q['option_c'],q['option_d'],
                        q['answer'],q['score'],q['topic']))
                    saved+=1
                except Exception as _e:
                    print(f"INSERT error: {_e}"); skipped+=1
            conn.commit(); conn.close()
            return jsonify({"success":True,"saved":saved,"skipped":skipped,"preview":questions[:3],"db_type":"pg" if (USE_PG and _HAS_PG) else "sqlite"})
        except Exception as e:
            return jsonify({"error":f"Файл уншихад алдаа: {str(e)}"}),500
    return render_template("admin_import.html",subjects=SUBJECTS,
        all_subjects=ALL_SUBJECTS,levels=LEVELS)

@app.route("/admin/ai-generate",methods=["GET","POST"])
@login_required
def admin_ai_generate():
    if request.method=="POST":
        import json as _json
        grade=int(request.form.get("grade",9)); subject=request.form.get("subject","Математик")
        topic=request.form.get("topic",""); lvl=request.form.get("level","Мэдлэг ойлголт")
        bloom=request.form.get("bloom","Мэдлэг"); q_type=request.form.get("q_type","Нэг сонголт")
        count=min(int(request.form.get("count",5)),20)
        api_key=os.environ.get("ANTHROPIC_API_KEY","")
        if not api_key:
            return jsonify({"error":"ANTHROPIC_API_KEY тохируулаагүй."}),400
        has_options=q_type in("Нэг сонголт","Олон сонголт")
        score=LEVEL_SCORE.get(lvl,1)
        is_math2 = subject in ['Математик','Алгебр','Геометр','Тригонометр']
        math_hint2 = ' Тоогоор бичигдсэн бодлого (тооцоолол, томьёо, геометр, тэгшитгэл) заавал оруул.' if is_math2 else ''
        prompt=(
             f'Монгол ЕБС {grade}-р анги {subject} хичээл.'
             + (f' Сэдэв: {topic}.' if topic else '')
             + f' Блюпринт: {lvl}, Блум: {bloom}, Төрөл: {q_type}.{math_hint2}\n'
             + f'{count} даалгавар үүсгэ. Монгол хэлээр. Зөвхөн JSON array:\n'
             + '[{"question":"...","option_a":"...","option_b":"...","option_c":"...","option_d":"...","answer":"А","topic":"' + (topic or subject) + '"}]'
             + (' Нээлттэй: option_a..d = null.' if not has_options else '')
        )
        try:
            client=anthropic.Anthropic(api_key=api_key)
            msg=client.messages.create(model="claude-sonnet-4-5",max_tokens=4000,
                messages=[{"role":"user","content":prompt}])
            raw_text=msg.content[0].text.strip()
            raw_text=re.sub(r'^```json\s*','',raw_text)
            raw_text=re.sub(r'^```\s*','',raw_text)
            raw_text=re.sub(r'\s*```$','',raw_text).strip()
            ai_qs=_json.loads(raw_text)
            if not isinstance(ai_qs,list): ai_qs=[ai_qs]
            conn=get_db(); saved=0
            for idx,q in enumerate(ai_qs):
                q_code=f"Q{grade}-{subject[:2]}-AI-{datetime.now().strftime('%f')}-{idx}"
                try:
                    conn.execute("""INSERT OR IGNORE INTO questions(q_code,grade,subject,level,bloom,q_type,
                       question,option_a,option_b,option_c,option_d,answer,score,topic)
                       VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                       (q_code,grade,subject,lvl,bloom,q_type,q.get('question',''),
                        q.get('option_a'),q.get('option_b'),q.get('option_c'),q.get('option_d'),
                        q.get('answer','А'),score,q.get('topic',topic)))
                    saved+=1
                except: pass
            conn.commit(); conn.close()
            return jsonify({"success":True,"saved":saved,"questions":ai_qs})
        except _json.JSONDecodeError as e:
            return jsonify({"error":f"JSON алдаа: {str(e)}"}),500
        except Exception as e:
            return jsonify({"error":str(e)}),500
    return render_template("admin_ai_generate.html",
        subjects=SUBJECTS, all_subjects=ALL_SUBJECTS,levels=LEVELS,blooms=BLOOM,q_types=Q_TYPES)

# ══ БЛЮПРИНТ ИМПОРТ (admin) ════════════════════════════════
@app.route("/admin/import-blueprint", methods=["GET", "POST"])
@login_required
def admin_import_blueprint():
    # GET — гараар бөглөх форм (PDF tab бас байгаа)
    if request.method == "GET":
        return render_template("admin_blueprint_form.html",
            all_subjects=ALL_SUBJECTS, saved=_get_blueprint_list())

    # POST — PDF-с задлах (Математикт ажиллана)
    f       = request.files.get("pdf")
    subject = request.form.get("subject", "Математик")
    grade   = int(request.form.get("grade", 9))
    msg_html = ""
    if not f or f.filename == "":
        msg_html = '<div class="msg-err">❌ PDF файл сонгоогүй байна</div>'
    else:
        try:
            from blueprint_data import parse_blueprint_pdf, _save_blueprint
            import tempfile, os as _os
            pdf_bytes = f.read()
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(pdf_bytes); tmp_path = tmp.name
            try:
                nj_list = parse_blueprint_pdf(tmp_path, subject, grade)
            finally:
                _os.unlink(tmp_path)
            if not nj_list:
                raise ValueError("PDF-с блюпринт олдсонгүй. 'Гараар оруулах' tab ашиглана уу.")
            if not _save_blueprint(subject, grade, nj_list):
                raise ValueError("Өгөгдлийн санд хадгалахад алдаа гарлаа")
            total = sum(len(s['shalguur']) for nj in nj_list for s in nj.get('surd', []))
            msg_html = (f'<div class="msg-ok">✅ Амжилттай! {subject} {grade}-р анги: '
                        f'{len(nj_list)} нэгж, {total} шалгуур</div>')
        except Exception as e:
            import traceback; traceback.print_exc()
            msg_html = f'<div class="msg-err">❌ {e}</div>'

    return render_template("admin_blueprint_form.html",
        all_subjects=ALL_SUBJECTS, saved=_get_blueprint_list(),
        pdf_msg=msg_html)

# ══ БАГШ ТАНД ══════════════════════════════════════════════
@app.route("/teacher")
def teacher_page():
    return render_template("teacher.html",
        subjects=SUBJECTS, all_subjects=ALL_SUBJECTS, levels=LEVELS, blooms=BLOOM,
        curriculum_grades=CURRICULUM.get("grades", {}),
        curriculum_periods=CURRICULUM.get("periods", ["40","80","90"]))

@app.route("/api/workplan", methods=["POST"])
def workplan():
    data    = request.json
    api_key = os.environ.get("ANTHROPIC_API_KEY","")
    if not api_key:
        return jsonify({"error":"API тохируулаагүй"}), 400
    if anthropic is None:
        return jsonify({"error": "anthropic суугаагүй"}), 500
    teacher = data.get("teacher",""); year = data.get("year","2025-2026")
    subject = data.get("subject",""); exp = data.get("exp","3")
    extra = data.get("extra",""); cats = data.get("cats",[])
    exp_labels = {"1":"1-3 жилийн туршлагатай залуу","2":"4-10 жилийн туршлагатай","3":"10+ жилийн ахлах"}
    cat_names = {"surgalt":"Сургалт","hugzhil":"Өөрийгөө хөгжүүлэх","ecej":"Асран хамгаалагчтай хамтрах",
                 "niigem":"Иргэд, олон нийттэй ажиллах","yos":"Ёс зүй","huuhded":"Хүүхэд хамгаалал"}
    selected_names = [cat_names.get(k,k) for k in cats]
    prompt = f"""Та Монгол ЕБС-ийн багшийн ажлын жилийн төлөвлөгөө боловсруулагч AI байна.

МЭДЭЭЛЭЛ:
- Багш: {teacher or 'тодорхойгүй'}
- Хичээлийн жил: {year}
- Хичээл/Анги: {subject or 'тодорхойгүй'}
- Туршлага: {exp_labels.get(exp,'10+ жил')}
{f'- Нэмэлт: {extra}' if extra else ''}
- Хэсгүүд: {', '.join(selected_names)}

ДААЛГАВАР:
Хэсэг бүрт 3-4 зорилт, зорилт бүрт үйл ажиллагаа, үр дүн, хугацаа бичиж JSON буцаана уу.

JSON формат:
{{
  {', '.join([f'"{k}": {{"goals": [{{"goal":"...", "actions":["...","..."], "result":"...", "time":"..."}}]}}' for k in cats[:2]])}
}}

Монгол хэлээр дэлгэрэнгүй. Зөвхөн JSON буцаана уу."""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(model="claude-sonnet-4-5", max_tokens=3000,
            messages=[{"role":"user","content":prompt}])
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"): raw = raw[4:]
            raw = raw.strip()
        import json as _j
        try:
            plan = _j.loads(raw)
            return jsonify({"plan": plan})
        except:
            return jsonify({"raw": raw})
    except Exception as e:
        import traceback; print(traceback.format_exc())
        return jsonify({"error": str(e)}), 500

@app.route("/certificate")
def certificate_page():
    return render_template("certificate.html", all_subjects=ALL_SUBJECTS)

@app.route("/api/certificate-ai", methods=["POST"])
def certificate_ai():
    import json as _j
    data = request.json
    api_key = os.environ.get("ANTHROPIC_API_KEY","")
    if not api_key:
        return jsonify({"error":"API тохируулаагүй"}), 400
    if anthropic is None:
        return jsonify({"error":"anthropic суугаагүй"}), 500
    cert_type = data.get("cert_type","award"); name = data.get("name","")
    reason = data.get("reason",""); school = data.get("school","Орхонтуул ЕБС")
    type_labels = {"award":"өргөмжлөл","cert":"сертификат","thanks":"талархал"}
    prompt = (
        school + "-н " + type_labels.get(cert_type,"өргөмжлөл") + " үүсгэж байна.\n"
        + ("Хүлээн авагч: " + name + "\n" if name else "")
        + ("Шалтгаан/Үйл ажиллагаа: " + reason + "\n" if reason else "")
        + "\nДараах JSON форматаар монгол хэлээр буцаа (зөвхөн JSON):\n"
        + '{"title":"ӨРГӨМЖЛӨЛ","subtitle":"Дэд гарчиг","value":"Товч тодорхойлолт"}'
    )
    try:
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(model="claude-sonnet-4-5", max_tokens=500,
            messages=[{"role":"user","content":prompt}])
        raw = msg.content[0].text.strip()
        if "```" in raw:
            for part in raw.split("```"):
                p2 = part.strip()
                if p2.startswith("json"): p2 = p2[4:].strip()
                if p2.startswith("{"): raw = p2; break
        s = raw.find("{"); e = raw.rfind("}")+1
        if s >= 0 and e > s: raw = raw[s:e]
        result = _j.loads(raw)
        return jsonify(result)
    except Exception as ex:
        import traceback; print(traceback.format_exc())
        return jsonify({"error": str(ex)}), 500

@app.route("/api/certificate", methods=["POST"])
def gen_cert():
    from certificate_gen import gen_certificate, gen_batch
    data = request.json
    cert_type = data.get("cert_type", "cert"); school = data.get("school", "Орхонтуул ЕБС")
    title = data.get("title", "СЕРТИФИКАТ"); subtitle = data.get("subtitle", "")
    date = data.get("date", ""); names = data.get("names", [])
    if not names:
        return jsonify({"error": "Нэр оруулаагүй байна"}), 400
    try:
        if len(names) == 1:
            pdf = gen_certificate(name=names[0]["name"], value=names[0].get("value",""),
                school=school, title=title, subtitle=subtitle, date=date, cert_type=cert_type)
        else:
            pdf = gen_batch(names, school=school, title=title, subtitle=subtitle, date=date, cert_type=cert_type)
        from flask import Response
        fname = "certificate.pdf" if len(names)==1 else "certificates_batch.pdf"
        return Response(pdf, mimetype="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={fname}"})
    except Exception as e:
        import traceback; print(traceback.format_exc())
        return jsonify({"error": str(e)}), 500

@app.route("/api/lesson-plan", methods=["POST"])
def lesson_plan():
    import json as _j2
    data     = request.json
    api_key  = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API тохируулаагүй"}), 400
    if anthropic is None:
        return jsonify({"error": "anthropic суугаагүй"}), 500
    manager  = str(data.get("manager_name") or ""); teacher = str(data.get("teacher_name") or "")
    subject  = str(data.get("subject") or "Математик"); topic = str(data.get("topic") or "Хичээлийн сэдэв")
    grade    = str(data.get("grade") or "9"); period = str(data.get("period") or "40")
    objectives = str(data.get("objectives") or "")
    grade_label = CURRICULUM.get("grades", {}).get(grade, {}).get("label", grade + "-р анги")
    p  = int(period); t1 = max(5, p // 8); t2 = p // 3 + p // 4; t3 = max(5, p - t1 - t2)
    stages_template = (
        '[{"name":"I.ЭХЛЭЛ","time":' + str(t1) + ',"purpose":"...","teacher_actions":"...","student_actions":"...","assessment":"..."},'
        '{"name":"II.СУДЛАЛ","time":' + str(t2) + ',"purpose":"...","teacher_actions":"...","student_actions":"...","assessment":"..."},'
        '{"name":"III.ДҮГНЭЛТ","time":' + str(t3) + ',"purpose":"...","teacher_actions":"...","student_actions":"...","assessment":"..."}]'
    )
    prompt = (
        "Монгол ЕБС-ийн " + grade_label + " " + subject + " хичээлийн "
        + topic + " сэдвийн " + period + " минутын ээлжит хичээлийн хөтөлбөр.\n"
        "Багш: " + (teacher or "тодорхойгүй") + " Менежер: " + (manager or "тодорхойгүй") + "\n"
        + ("Зорилт: " + objectives + "\n" if objectives else "")
        + "\nМОНГОЛ ХЭЛЭЭР дэлгэрэнгүй бөглөж зөвхөн JSON буцаа:\n"
        + '{"header":{"subject":"' + subject + '","topic":"' + topic + '","grade":"' + grade_label + '","period":"' + period + 'мин","teacher":"' + teacher + '","manager":"' + manager + '"},'
        + '"objectives":{"A":"' + topic + ' үндсэн ойлголтыг мэддэг болно","B":"' + topic + ' тайлбарлаж чадна","C":"' + topic + ' амьдралд хэрэглэнэ"},'
        + '"design":{"method":"Bloom таксономи","tools":"Сурах бичиг, самбар","engagement":"Бүлгийн ажил"},'
        + '"stages":' + stages_template + ','
        + '"differentiation":{"support":"Нэмэлт тайлбар","advanced":"Нэмэлт бодлого"},'
        + '"homework":"' + topic + '-той холбоотой гэрийн даалгавар"}'
    )
    try:
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(model="claude-sonnet-4-5", max_tokens=4000,
            messages=[{"role": "user", "content": prompt}])
        raw = msg.content[0].text.strip()
        if "```" in raw:
            for part in raw.split("```"):
                p2 = part.strip()
                if p2.startswith("json"): p2 = p2[4:].strip()
                if p2.startswith("{"): raw = p2; break
        s = raw.find("{"); e = raw.rfind("}") + 1
        if s >= 0 and e > s: raw = raw[s:e]
        plan = _j2.loads(raw)
        if "header" in plan:
            plan["header"].update({"subject":subject,"topic":topic,"grade":grade_label,"period":period+"мин","teacher":teacher,"manager":manager})
        return jsonify({"plan": plan})
    except Exception as ex:
        import traceback; print("LESSON ERROR:", traceback.format_exc())
        return jsonify({"error": str(ex)}), 500

@app.route("/api/teacher-ai", methods=["POST"])
def teacher_ai():
    data    = request.json
    prompt  = data.get("prompt","")
    api_key = os.environ.get("ANTHROPIC_API_KEY","")
    if not api_key:
        return jsonify({"error":"ANTHROPIC_API_KEY тохируулаагүй."}), 400
    if anthropic is None:
        return jsonify({"error": "anthropic суугаагүй"}), 500
    if not prompt:
        return jsonify({"error":"Prompt хоосон байна"}), 400
    try:
        client  = anthropic.Anthropic(api_key=api_key)
        msg     = client.messages.create(model="claude-sonnet-4-5", max_tokens=3000,
            messages=[{"role":"user","content": prompt}])
        result  = msg.content[0].text.strip()
        return jsonify({"result": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ══ ЧАТБОТ ════════════════════════════════════════════════
@app.route("/api/chat", methods=["POST"])
def chatbot():
    data     = request.json
    messages = data.get("messages", [])
    api_key  = os.environ.get("ANTHROPIC_API_KEY","")
    if not api_key:
        return jsonify({"error":"API тохируулаагүй."}), 400
    if anthropic is None:
        return jsonify({"error": "anthropic суугаагүй"}), 500
    if not messages:
        return jsonify({"error":"Мессеж хоосон"}), 400
    try:
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(model="claude-sonnet-4-5", max_tokens=2000,
            system="Та Орхонтуул ЕБС-ийн ухаалаг туслах AI. Багш, сурагчид Монгол хэлээр дэлгэрэнгүй, найрсаг хариулна.",
            messages=messages)
        return jsonify({"reply": msg.content[0].text})
    except Exception as e:
        import traceback; print("CHAT ERROR:", traceback.format_exc())
        return jsonify({"error": str(e)}), 500

@app.route("/teacher-login", methods=["GET", "POST"])
def teacher_login_page():
    if request.method == "POST":
        pw = request.form.get("password", "")
        if pw == TEACHER_PASSWORD:
            session.permanent = True
            session["teacher_logged_in"] = True
            return redirect(url_for("questions_page"))
        return render_template("teacher_login.html", error="Нууц үг буруу байна")
    return render_template("teacher_login.html", error=None)

@app.route("/teacher-logout")
def teacher_logout():
    session.pop("teacher_logged_in", None)
    return redirect(url_for("questions_page"))

@app.route("/api/upload", methods=["POST"])
def api_upload():
    if not session.get("teacher_logged_in") and not session.get("admin"):
        return jsonify({"error": "teacher_login_required", "redirect": "/teacher-login"}), 401
    from file_importer import extract_from_pdf, extract_from_docx, parse_raw_text
    f = request.files.get("file")
    grade = int(request.form.get("grade", 9))
    subject = request.form.get("subject", "Математик")
    lvl = request.form.get("default_level", "auto")
    if not f or f.filename == "":
        return jsonify({"error": "Файл сонгоогүй байна"}), 400
    file_bytes = f.read(); fname = f.filename.lower()
    try:
        if fname.endswith(".pdf"):
            raw = extract_from_pdf(file_bytes)
        elif fname.endswith((".docx", ".doc")):
            raw = extract_from_docx(file_bytes)
        elif fname.endswith(".txt"):
            raw = file_bytes.decode("utf-8", errors="ignore")
        elif fname.endswith((".jpg", ".jpeg", ".png")):
            import base64, anthropic
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
            if not api_key:
                return jsonify({"error": "API key байхгүй"}), 400
            b64 = base64.b64encode(file_bytes).decode()
            mt = "image/jpeg" if fname.endswith((".jpg", ".jpeg")) else "image/png"
            cl = anthropic.Anthropic(api_key=api_key)
            msg = cl.messages.create(model="claude-sonnet-4-5", max_tokens=3000,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": mt, "data": b64}},
                    {"type": "text", "text": "Даалгаваруудыг задлан дугаарлан жагсаа."}
                ]}])
            raw = msg.content[0].text
        else:
            return jsonify({"error": "PDF, Word, JPG эсвэл TXT файл оруулна уу"}), 400
        questions = parse_raw_text(raw, grade, subject, lvl)
        if not questions:
            return jsonify({"error": "Даалгавар илрүүлж чадсангүй"}), 400
        init_db()
        conn = get_db(); saved = skipped = 0
        for q in questions:
            try:
                conn.execute("""INSERT INTO questions(q_code,grade,subject,level,bloom,q_type,
                    question,option_a,option_b,option_c,option_d,answer,score,topic)
                    VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (q['q_code'],q['grade'],q['subject'],q.get('level',lvl),q['bloom'],
                    q['q_type'],q['question'],q['option_a'],q['option_b'],
                    q['option_c'],q['option_d'],q['answer'],q['score'],q['topic']))
                saved += 1
            except Exception as err:
                print(f"Upload INSERT error: {err}"); skipped += 1
        conn.commit(); conn.close()
        return jsonify({"success": True, "saved": saved, "skipped": skipped,
                        "db_type": "pg" if (USE_PG and _HAS_PG) else "sqlite"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ══ START ══════════════════════════════════════════════════
try:
    init_db()
    print("✅ DB tables ready")
except Exception as _ie:
    print(f"⚠️ init_db: {_ie}")

try:
    if _HAS_BLUEPRINT:
        init_blueprint_table()
        print("✅ Blueprint table ready")
except Exception as _be:
    print(f"⚠️ blueprint table: {_be}")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
